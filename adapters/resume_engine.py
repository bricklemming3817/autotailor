"""Stub resume engine that returns PDF and DOCX bytes.

Contract:

def generate_both(profile: dict, jd_url: str) -> dict:
    Returns a dict with keys: pdf (bytes), docx (bytes), filenames{pdf, docx}, coverage.
"""

from __future__ import annotations

from datetime import datetime
from urllib.parse import urlparse

from weasyprint import HTML
from docx import Document


def _build_text(profile: dict, jd_url: str) -> str:
    name = (profile.get("full_name") or "").strip()
    lines = []
    if name:
        lines.append(name)
    contact = []
    for key in ("city", "email", "phone", "linkedin", "github"):
        val = (profile.get(key) or "").strip()
        if val:
            contact.append(val)
    if contact:
        lines.append(" | ".join(contact))
    lines.append("")
    about = (profile.get("about") or "").strip()
    if about:
        lines.append("Summary")
        lines.append(about)
        lines.append("")
    lines.append(f"(Stub) Tailored for: {jd_url}")
    lines.append("")
    lines.append("Experience")
    lines.append("• Contributed to team projects with clear communication and documentation.")
    lines.append("• Applied standard problem-solving and data handling techniques where appropriate.")
    return "\n".join(lines)


def _to_pdf_bytes(text: str) -> bytes:
    html = f"""
    <html>
      <head>
        <meta charset='utf-8'>
        <style>
          body {{ font-family: -apple-system, Segoe UI, Roboto, sans-serif; font-size: 12pt; }}
          pre {{ white-space: pre-wrap; }}
        </style>
      </head>
      <body>
        <pre>{text}</pre>
      </body>
    </html>
    """
    return HTML(string=html).write_pdf()


def _to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        if line.startswith("• "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            # Keeping bullets simple in stub; no fabricated claims
        else:
            doc.add_paragraph(line)
    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_both(profile: dict, jd_url: str) -> dict:
    """Generate stub resume PDF and DOCX bytes with filenames and coverage.

    Returns:
      {
        "pdf": bytes,
        "docx": bytes,
        "filenames": {
          "pdf": "Resume_<Name_underscored>_<host>_<YYYYMMDD>.pdf",
          "docx":"Resume_<Name_underscored>_<host>_<YYYYMMDD>.docx"
        },
        "coverage": {"score": 0.8, "hits": ["sql"], "misses": ["dbt"]}
      }
    """
    text = _build_text(profile, jd_url)
    pdf_bytes = _to_pdf_bytes(text)
    docx_bytes = _to_docx_bytes(text)

    name = (profile.get("full_name") or "").strip() or "Resume"
    name_safe = "_".join(name.split())
    host = urlparse(jd_url).hostname or "job"
    today = datetime.utcnow().strftime("%Y%m%d")
    filenames = {
        "pdf": f"Resume_{name_safe}_{host}_{today}.pdf",
        "docx": f"Resume_{name_safe}_{host}_{today}.docx",
    }

    coverage = {"score": 0.8, "hits": ["sql"], "misses": ["dbt"]}

    return {"pdf": pdf_bytes, "docx": docx_bytes, "filenames": filenames, "coverage": coverage}
